from io import BytesIO

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
)


def export_csv(df):
    return df.to_csv(index=False).encode("utf-8-sig")


def export_excel(df):
    output = BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Applications")

        total = len(df)
        interview = int((df["Status"] == "Interview").sum())
        offer = int((df["Status"] == "Offer").sum())

        summary = pd.DataFrame({
            "Metric": [
                "Total", "Saved", "Applied", "Assessment",
                "Interview", "Offer", "Rejected",
                "Interview Rate", "Offer Rate"
            ],
            "Value": [
                total,
                int((df["Status"] == "Saved").sum()),
                int((df["Status"] == "Applied").sum()),
                int((df["Status"] == "Assessment").sum()),
                interview,
                offer,
                int((df["Status"] == "Rejected").sum()),
                f"{(interview / total * 100) if total else 0:.1f}%",
                f"{(offer / total * 100) if total else 0:.1f}%",
            ],
        })
        summary.to_excel(writer, index=False, sheet_name="Summary")

    output.seek(0)
    wb = load_workbook(output)

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)

    for ws in wb.worksheets:
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")

        for column_cells in ws.columns:
            max_len = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in column_cells
            )
            width = min(max(max_len + 2, 12), 45)
            ws.column_dimensions[get_column_letter(column_cells[0].column)].width = width

        ws.freeze_panes = "A2"

    ws = wb["Applications"]
    status_colors = {
        "Saved": "ADD8E6",
        "Applied": "FFA500",
        "Assessment": "FFF2CC",
        "Interview": "87CEEB",
        "Offer": "90EE90",
        "Rejected": "FF7F7F",
    }

    status_col = None
    for cell in ws[1]:
        if cell.value == "Status":
            status_col = cell.column
            break

    if status_col:
        for row in range(2, ws.max_row + 1):
            status = ws.cell(row, status_col).value
            if status in status_colors:
                fill = PatternFill(
                    "solid",
                    fgColor=status_colors[status],
                )
                ws.cell(row, status_col).fill = fill

    final = BytesIO()
    wb.save(final)
    return final.getvalue()


def export_word(df):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    header = section.header.paragraphs[0]
    header.text = "Smart Job Application Tracker"
    header.runs[0].font.size = Pt(10)

    footer = section.footer.paragraphs[0]
    footer.text = "Generated report"
    footer.runs[0].font.size = Pt(9)

    title = doc.add_heading("Job Applications Report", level=1)
    title.runs[0].font.size = Pt(20)

    doc.add_paragraph(
        f"Total applications in this export: {len(df)}"
    )

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(8)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def export_pdf(df):
    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=24,
        leftMargin=24,
        topMargin=28,
        bottomMargin=28,
    )

    styles = getSampleStyleSheet()
    story = [
        Paragraph("Smart Job Application Tracker", styles["Title"]),
        Paragraph(
            f"Job Applications Report — {len(df)} application(s)",
            styles["Normal"],
        ),
        Spacer(1, 12),
    ]

    data = [list(df.columns)]
    for _, row in df.iterrows():
        data.append([str(value) for value in row])

    # Keep the PDF readable even with long URLs.
    data = [
        [
            Paragraph(str(cell), styles["BodyText"])
            for cell in row
        ]
        for row in data
    ]

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F5F7FA")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))

    story.append(table)
    doc.build(story)
    return output.getvalue()
